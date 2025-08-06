"""Document processing utilities for the knowledge base."""

import os
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple

try:
    from pypdf import PdfReader
except ImportError:
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        PdfReader = None

try:
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup
except ImportError:
    ebooklib = None
    epub = None
    BeautifulSoup = None

import docx

from src.utils.logging import setup_logger

# Set up module logger
logger = setup_logger("document_processor")


class DocumentProcessor:
    """Processes various document formats for knowledge extraction."""
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> Tuple[str, Dict]:
        """
        Extract text and metadata from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Tuple of (extracted text, metadata dictionary)
        """
        if PdfReader is None:
            logger.error("PDF processing not available - install pypdf or PyPDF2")
            return "", {"error": "PDF processing not available", "format": "pdf"}
        
        try:
            reader = PdfReader(file_path)
            text = ""
            
            # Extract text from each page
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():  # Only add non-empty pages
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
            
            # Clean up text
            text = DocumentProcessor._clean_text(text)
            
            # Extract basic metadata
            metadata = {
                "title": os.path.basename(file_path),
                "pages": len(reader.pages),
                "format": "pdf",
                "file_path": file_path
            }
            
            # Try to get document info if available
            if hasattr(reader, 'metadata') and reader.metadata:
                try:
                    metadata.update({
                        "title": reader.metadata.get("/Title", metadata["title"]) or metadata["title"],
                        "author": reader.metadata.get("/Author", "Unknown") or "Unknown",
                        "creation_date": str(reader.metadata.get("/CreationDate", "Unknown")) or "Unknown"
                    })
                except Exception as e:
                    logger.warning(f"Error extracting PDF metadata: {e}")
            
            logger.info(f"Processed PDF: {file_path} ({len(reader.pages)} pages, {len(text)} chars)")
            return text, metadata
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {str(e)}")
            return "", {"error": str(e), "format": "pdf", "file_path": file_path}
    
    @staticmethod
    def extract_from_epub(file_path: str) -> Tuple[str, Dict]:
        """
        Extract text and metadata from an EPUB file.
        
        Args:
            file_path: Path to the EPUB file
            
        Returns:
            Tuple of (extracted text, metadata dictionary)
        """
        if ebooklib is None or epub is None or BeautifulSoup is None:
            logger.error("EPUB processing not available - install ebooklib and beautifulsoup4")
            return "", {"error": "EPUB processing not available", "format": "epub"}
        
        try:
            book = epub.read_epub(file_path)
            text = ""
            
            # Extract metadata
            metadata = {
                "title": os.path.basename(file_path),
                "format": "epub",
                "file_path": file_path
            }
            
            # Update with book metadata if available
            try:
                if book.get_metadata('DC', 'title'):
                    metadata["title"] = book.get_metadata('DC', 'title')[0][0]
                    
                if book.get_metadata('DC', 'creator'):
                    metadata["author"] = book.get_metadata('DC', 'creator')[0][0]
                    
                if book.get_metadata('DC', 'language'):
                    metadata["language"] = book.get_metadata('DC', 'language')[0][0]
            except Exception as e:
                logger.warning(f"Error extracting EPUB metadata: {e}")
            
            # Extract content from each document
            chapter_count = 0
            for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
                try:
                    content = item.get_content()
                    soup = BeautifulSoup(content, 'html.parser')
                    chapter_text = soup.get_text()
                    
                    if chapter_text.strip():  # Only add non-empty chapters
                        text += f"\n--- Chapter {chapter_count + 1} ---\n{chapter_text}\n"
                        chapter_count += 1
                except Exception as e:
                    logger.warning(f"Error processing EPUB chapter: {e}")
            
            # Clean up text
            text = DocumentProcessor._clean_text(text)
            metadata["chapters"] = chapter_count
            
            logger.info(f"Processed EPUB: {file_path} ({chapter_count} chapters, {len(text)} chars)")
            return text, metadata
            
        except Exception as e:
            logger.error(f"Error processing EPUB {file_path}: {str(e)}")
            return "", {"error": str(e), "format": "epub", "file_path": file_path}
            return "", {"error": str(e), "format": "epub"}
    
    @staticmethod
    def extract_from_txt(file_path: str) -> Tuple[str, Dict]:
        """
        Extract text and basic metadata from a text file.
        
        Args:
            file_path: Path to the text file
            
        Returns:
            Tuple of (extracted text, metadata dictionary)
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            text = ""
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if not text:
                raise ValueError("Could not decode file with any supported encoding")
            
            # Clean up text
            text = DocumentProcessor._clean_text(text)
            
            metadata = {
                "title": os.path.basename(file_path),
                "format": "txt",
                "file_path": file_path,
                "encoding": encoding
            }
            
            logger.info(f"Processed text file: {file_path} ({len(text)} chars)")
            return text, metadata
            
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {str(e)}")
            return "", {"error": str(e), "format": "txt", "file_path": file_path}
    
    @staticmethod
    def extract_from_docx(file_path: str) -> Tuple[str, Dict]:
        """
        Extract text and metadata from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Tuple of (extracted text, metadata dictionary)
        """
        try:
            doc = docx.Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Clean up text
            text = DocumentProcessor._clean_text(text)
            
            # Extract metadata
            metadata = {
                "title": os.path.basename(file_path),
                "format": "docx",
                "file_path": file_path,
                "paragraphs": len(doc.paragraphs)
            }
            
            # Try to get document properties
            try:
                core_props = doc.core_properties
                if core_props.title:
                    metadata["title"] = core_props.title
                if core_props.author:
                    metadata["author"] = core_props.author
                if core_props.created:
                    metadata["creation_date"] = str(core_props.created)
            except Exception as e:
                logger.warning(f"Error extracting DOCX metadata: {e}")
            
            logger.info(f"Processed DOCX: {file_path} ({len(doc.paragraphs)} paragraphs, {len(text)} chars)")
            return text, metadata
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {str(e)}")
            return "", {"error": str(e), "format": "docx", "file_path": file_path}
    
    @staticmethod
    def _clean_text(text: str) -> str:
        """
        Clean and normalize extracted text.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)  # Multiple newlines to double
        text = re.sub(r' +', ' ', text)  # Multiple spaces to single
        text = re.sub(r'\t+', ' ', text)  # Tabs to spaces
        
        # Remove page breaks and form feeds
        text = text.replace('\f', '\n')
        text = text.replace('\r', '\n')
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    @classmethod
    def process_file(cls, file_path: str) -> Tuple[str, Dict]:
        """
        Process a file based on its extension.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Tuple of (extracted text, metadata dictionary)
        """
        file_path = str(file_path)  # Ensure string type
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.pdf':
            return cls.extract_from_pdf(file_path)
        elif ext == '.epub':
            return cls.extract_from_epub(file_path)
        elif ext == '.txt':
            return cls.extract_from_txt(file_path)
        elif ext == '.docx':
            return cls.extract_from_docx(file_path)
        else:
            logger.warning(f"Unsupported file format: {ext} ({file_path})")
            return "", {"error": "Unsupported format", "format": ext[1:], "file_path": file_path}
    
    @classmethod
    def get_supported_extensions(cls) -> List[str]:
        """Get list of supported file extensions."""
        return ['.pdf', '.epub', '.txt', '.docx']
    
    @classmethod
    def process_directory(cls, directory_path: str) -> List[Tuple[str, Dict, str]]:
        """
        Process all supported files in a directory.
        
        Args:
            directory_path: Path to the directory
            
        Returns:
            List of tuples (text, metadata, file_path)
        """
        supported_extensions = cls.get_supported_extensions()
        results = []
        
        try:
            directory = Path(directory_path)
            if not directory.exists() or not directory.is_dir():
                logger.error(f"Directory not found: {directory_path}")
                return results
            
            # Recursively find all supported files
            for file_path in directory.rglob("*"):
                if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                    logger.info(f"Processing: {file_path}")
                    text, metadata = cls.process_file(file_path)
                    if text and not metadata.get("error"):  # Only add if successful
                        results.append((text, metadata, str(file_path)))
                    elif metadata.get("error"):
                        logger.warning(f"Skipping {file_path} due to error: {metadata['error']}")
            
            logger.info(f"Successfully processed {len(results)} documents from {directory_path}")
            return results
            
        except Exception as e:
            logger.error(f"Error processing directory {directory_path}: {str(e)}")
            return results